from __future__ import annotations

import ast
from datetime import datetime
from pathlib import Path

from docx import Document

from src.svm_model import SVMClassifier, load_yolo_class_names


def _parse_names_from_yaml(data_yaml: Path) -> list[str]:
    names_line = None
    for line in data_yaml.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if s.startswith("names:"):
            names_line = s.split("names:", 1)[1].strip()
            break
    if not names_line:
        raise ValueError("Cannot find names in data.yaml")
    names = ast.literal_eval(names_line)
    if not isinstance(names, list):
        raise ValueError("names in data.yaml must be a list")
    return [str(x) for x in names]


def _collect_label_stats(labels_dir: Path, class_names: list[str]) -> tuple[dict[str, int], int]:
    per_class = {name: 0 for name in class_names}
    invalid_files = 0

    for txt_path in sorted(labels_dir.glob("*.txt")):
        try:
            first = ""
            with txt_path.open("r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line:
                        first = line
                        break
            if not first:
                invalid_files += 1
                continue
            parts = first.split()
            class_id = int(parts[0])
            if class_id < 0 or class_id >= len(class_names):
                invalid_files += 1
                continue
            per_class[class_names[class_id]] += 1
        except (OSError, ValueError, IndexError):
            invalid_files += 1

    return per_class, invalid_files


def _fmt(value: float) -> str:
    return f"{value:.4f}"


def build_report(root: Path) -> Path:
    dataset_dir = root / "dataset"
    data_yaml = dataset_dir / "data.yaml"
    images_dir = dataset_dir / "train" / "images"
    labels_dir = dataset_dir / "train" / "labels"

    class_names = _parse_names_from_yaml(data_yaml)
    per_class_counts, invalid_labels = _collect_label_stats(labels_dir, class_names)
    n_images = len(list(images_dir.glob("*")))
    n_labels = len(list(labels_dir.glob("*.txt")))

    classifier = SVMClassifier(kernel="rbf", C=1.0, gamma="scale", image_size=(64, 64))
    stats = classifier.fit(
        images_dir=images_dir,
        annotations_dir=labels_dir,
        class_names=class_names,
        test_size=0.2,
    )

    out_path = root / "Assignment2_Report_Final_FromTemplate.docx"

    doc = Document()
    doc.add_heading("COMP4423 - Assignment 2", level=0)
    doc.add_paragraph("Student Name (Student ID)")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report documents an end-to-end campus vegetation recognition system built with "
        "traditional machine learning only. The implementation is coherent with the submitted code "
        "(main.py, src/svm_model.py, and app.py) and includes dataset organization, training, evaluation, "
        "error analysis, and a local application for prediction."
    )

    doc.add_heading("2. Method", level=1)
    doc.add_paragraph("In your report, these questions should be answered:")

    doc.add_heading("Q1. How do you design and test the program?", level=2)
    doc.add_paragraph(
        "Program design follows a reproducible pipeline: parse YOLO labels, extract handcrafted image "
        "features, train an SVM classifier, and evaluate on a held-out split. The project keeps clear entry "
        "points: main.py for training/evaluation and app.py for local inference."
    )
    doc.add_paragraph(
        f"Dataset summary: {len(class_names)} classes, {n_images} training images, {n_labels} label files."
    )

    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "Class"
    t1.rows[0].cells[1].text = "Image Count (by first-object label)"
    for cls_name in class_names:
        row = t1.add_row().cells
        row[0].text = cls_name
        row[1].text = str(per_class_counts.get(cls_name, 0))

    doc.add_paragraph(
        "Testing and sanity checks: path validation, label parsing validation, invalid-label counting, "
        f"and split-safety checks. Invalid/corrupted label files detected: {invalid_labels}."
    )

    doc.add_heading("Q2. How do you ensure the robustness of the program in a real scenario?", level=2)
    doc.add_paragraph(
        "The code reduces background shortcut learning by combining global, bounding-box-focused, and "
        "polygon-masked features. It also performs group-aware splitting when possible to reduce near-duplicate "
        "leakage between train and validation sets."
    )
    doc.add_paragraph(
        "At inference time, the app provides confidence and top-k outputs, optional automatic target attention, "
        "image enhancement, and manual ROI selection to improve reliability on field photos with complex backgrounds."
    )

    doc.add_heading("Q3. What problems do you find and how do you solve them?", level=2)
    doc.add_paragraph(
        "Problem 1: Background bias and overfitting to context.\n"
        "Solution: Added object-focused region and polygon-masked features; reduced global-context weight.\n\n"
        "Problem 2: Similar-looking species and illumination/viewpoint shifts cause confusion.\n"
        "Solution: Added richer metric reporting (macro/weighted/per-class), analyzed confusion matrix, and "
        "introduced inference-time attention fusion and ROI tools.\n\n"
        "Problem 3: UI dependency compatibility issues.\n"
        "Solution: Pinned compatible versions for streamlit and streamlit-drawable-canvas, and updated UI API calls."
    )

    doc.add_heading("Q4. How do you use GenAI to assist you in the implementation?", level=2)
    doc.add_paragraph(
        "GenAI was used to accelerate code drafting, refactoring ideas, and debugging suggestions, especially for "
        "dataset-format migration (to YOLO-only), evaluation output enrichment, and app interaction improvements."
    )

    doc.add_heading("Q5. How does GenAI understand and solve the tasks?", level=2)
    doc.add_paragraph(
        "GenAI first interprets task constraints (traditional ML only, local app required, report template questions), "
        "then proposes incremental code modifications and validates behavior through runtime feedback. The final workflow "
        "is iterative: propose -> implement -> test -> analyze errors -> patch."
    )

    doc.add_heading("Q6. What are the limitations and areas for improvement in GenAI's solution?", level=2)
    doc.add_paragraph(
        "Limitations observed: version assumptions can be wrong, and generated solutions may initially miss environment-specific "
        "constraints or edge cases. Improvements: stronger automatic environment checks, stricter compatibility planning before "
        "feature integration, and more explicit evaluation protocols for real-world generalization."
    )

    doc.add_heading("3. Results and Discussion", level=1)
    doc.add_paragraph(
        "Latest reproducible run (main.py default settings):"
    )

    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Metric"
    t2.rows[0].cells[1].text = "Value"
    metric_rows = [
        ("Train Accuracy", _fmt(float(stats["train_accuracy"]))),
        ("Validation Accuracy", _fmt(float(stats["val_accuracy"]))),
        ("Validation Macro Precision", _fmt(float(stats["val_macro_precision"]))),
        ("Validation Macro Recall", _fmt(float(stats["val_macro_recall"]))),
        ("Validation Macro F1", _fmt(float(stats["val_macro_f1"]))),
        ("Validation Weighted F1", _fmt(float(stats["val_weighted_f1"]))),
        ("Number of Classes", str(stats["num_classes"])),
        ("Number of Images", str(stats["num_examples"])),
        ("Number of Groups", str(stats["num_groups"])),
        ("Used Group Split", str(stats["used_group_split"])),
    ]
    for k, v in metric_rows:
        row = t2.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph(
        "Feature strategy comparison from the same run: "
        f"mixed val_acc={_fmt(float(stats['comparison']['mixed']['val_accuracy']))}, "
        f"mixed val_macro_f1={_fmt(float(stats['comparison']['mixed']['val_macro_f1']))}; "
        f"masked_only val_acc={_fmt(float(stats['comparison']['masked_only']['val_accuracy']))}, "
        f"masked_only val_macro_f1={_fmt(float(stats['comparison']['masked_only']['val_macro_f1']))}."
    )

    doc.add_paragraph(
        "Error analysis summary: difficult cases are mainly caused by similar morphology between species, "
        "background interference, and lighting/viewpoint shifts. These failures are consistent with limitations "
        "of handcrafted grayscale/shape-sensitive features under large appearance changes."
    )

    doc.add_heading("4. Application", level=1)
    doc.add_paragraph(
        "A local Streamlit application is implemented in app.py. It supports uploading or selecting an image "
        "and outputs predicted class, confidence score, and top-k candidates. It also includes optional enhancement, "
        "automatic attention, and manual ROI for practical field usage."
    )
    doc.add_paragraph("Launch command: python -m streamlit run app.py")

    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "The assignment requirements are fulfilled with a traditional ML approach, reproducible training code, "
        "detailed evaluation, and a usable local application. Future work can improve generalization by adding more "
        "diverse collection conditions, stronger handcrafted features, and more rigorous cross-location validation."
    )

    doc.add_paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parent
    output = build_report(project_root)
    print(f"saved: {output.name}")
